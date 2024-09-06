import os
import json
import argparse
from docx import Document
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from datetime import datetime
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement

def append_review_to_docx(docx_file, image_path, digest, revision_text, revision_edits, metadata):
    """Add a table with an image on the left and combined text on the right to the .docx file."""
    doc = Document(docx_file) if os.path.exists(docx_file) else Document()
    doc.add_page_break()

    # Add a table with 2 rows and 3 columns
    table = doc.add_table(rows=2, cols=3)

    # Set the table to take the whole page width
    table_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    for cell in table.columns[0].cells:
        cell.width = table_width / 3
    for cell in table.columns[1].cells:
        cell.width = table_width / 3
    for cell in table.columns[2].cells:
        cell.width = table_width / 3

    # First row, first cell: Add the image
    cell1 = table.cell(0, 0)
    run = cell1.paragraphs[0].add_run()
    run.add_picture(image_path, width=Inches(1.5))  # Adjust image size

    # First row, second cell: Add the digest text
    cell2 = table.cell(0, 1)
    cell2.text = digest

    # First row, third cell: Add the revision text
    cell3 = table.cell(0, 2)
    cell3.text = revision_text

    # Second row, first cell: Add metadata
    cell4 = table.cell(1, 0)
    cell4.text = metadata

    # Second row, second and third cells: Merge and add revision edits
    cell5 = table.cell(1, 1)
    cell6 = table.cell(1, 2)
    cell5.merge(cell6)
    cell5.text = revision_edits

    # Set font size for digest and revision text to 8 (smaller)
    for cell in [table.cell(0, 1), table.cell(0, 2)]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)  # Smaller font size

    # Set font size for metadata to 6
    for paragraph in table.cell(1, 0).paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(6)

    doc.save(docx_file)

def digest_to_text(word_objects):
    """
    Converts the processed word objects into a single piece of text.
    Line breaks should be represented as newlines, and words should be separated by spaces.
    """
    text_lines = []
    current_line = []

    for word_obj in word_objects:
        word = word_obj['word']
        detected_break = word_obj['detected_break']

        # Add the word to the current line
        current_line.append(word)

        # If the detected break is a newline (3 or 5), flush the current line
        if detected_break in [3, 5]:
            text_lines.append(' '.join(current_line))
            current_line = []  # Start a new line

    # Add any remaining words that didn't trigger a newline
    if current_line:
        text_lines.append(' '.join(current_line))

    # Join all the lines with newlines and return the final text
    return '\n'.join(text_lines)

def load_revision(revision_path):
    with open(revision_path, 'r', encoding='utf-8') as f:
        revision_data = json.load(f)
    return revision_data['text'], revision_data['edits']

def main():
    parser = argparse.ArgumentParser(description="Populate a .docx with digests, revisions, and images for manual review.")
    parser.add_argument('docx_file', help='Path to the .docx file to populate or create')
    parser.add_argument('image_file', help='Path to the image file')
    parser.add_argument('digest_file', help='Path to the digest JSON file')
    parser.add_argument('revision_file', help='Path to the revision JSON file')
    args = parser.parse_args()

    load_dotenv()

    docx_file = args.docx_file
    image_path = args.image_file
    digest_path = args.digest_file
    revision_path = args.revision_file

    # Check if both digest and image files exist
    if not os.path.exists(digest_path) or not os.path.exists(image_path):
        print(f"Error: Both digest file '{digest_path}' and image file '{image_path}' must exist. Exiting.")
        exit(1)

    # Load digest content
    with open(digest_path, 'r', encoding='utf-8') as f:
        digest_data = json.load(f)

    # Add the image and all text entries to the .docx file as a single page
    digest_text = digest_to_text(digest_data)
    revision_text, revision_edits = load_revision(revision_path)
    metadata = f"Image: {os.path.basename(image_path)}\nDigest: {os.path.basename(digest_path)}\nRevision: {os.path.basename(revision_path)}\nTime: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    append_review_to_docx(docx_file, image_path, digest_text, revision_text, revision_edits, metadata)

    print(f'Processed digest {digest_path} with image {image_path}')

if __name__ == '__main__':
    main()
